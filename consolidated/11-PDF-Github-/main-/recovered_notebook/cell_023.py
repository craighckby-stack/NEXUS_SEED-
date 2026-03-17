```python
from google.colab import files
import shutil
import os
from docx import Document

# Step 1: Define File Structure
os.makedirs("MIA-OS", exist_ok=True)

# Core scripts
scripts = {
    "bootstrap.sh": """#!/bin/bash
set -e
apt update && apt install -y python3 python3-pip nftables git
echo 'MIA OS Bootstrap Complete.'""",
    "install_mia.sh": """#!/bin/bash
set -e
mkdir -p /opt/mia
echo 'MIA OS Installed Successfully.'""",
    "firewall.nft": """table inet mia_firewall {
    chain input {
        type filter hook input priority 0; policy drop;
    }
}""",
    "mia.service": """[Unit]
Description=MIA OS Core Service
After=network.target
[Service]
ExecStart=/usr/bin/python3 /opt/mia/core_service.py
Restart=always
[Install]
WantedBy=multi-user.target""",
    "config.yaml": """firewall:
    enabled: true
    rules_file: /etc/mia/firewall.nft""",
    "core_service.py": """import time
while True:
    print('MIA OS Running...')
    time.sleep(60)
"""
}

# Step 2: Create and Save Files
for filename, content in scripts.items():
    with open(f"MIA-OS/{filename}", "w") as f:
        f.write(content)

# Step 3: Package Files
shutil.make_archive("MIA-OS-Package", 'gztar', "MIA-OS")

# Step 4: Generate Documentation
doc = Document()
doc.add_heading('MIA OS Test Version - Documentation', 0)
doc.add_paragraph('This document contains installation steps and details for MIA OS.')
doc.add_heading('Installation Steps:', level=1)
doc.add_paragraph('1. Extract the package: `tar -xzf MIA-OS-Package.tar.gz`')
doc.add_paragraph('2. Run the bootstrap script: `bash bootstrap.sh`')
doc.add_paragraph('3. Install MIA OS: `bash install_mia.sh`')
doc.add_paragraph('4. Start the service: `systemctl start mia.service`')
doc.save("MIA-OS-Documentation.docx")

# Step 5: Output Files for Download
files.download("MIA-OS-Package.tar.gz")
files.download("MIA-OS-Documentation.docx")

# Output: 
# MIA-OS-Package.tar.gz and MIA-OS-Documentation.docx files are downloaded
```