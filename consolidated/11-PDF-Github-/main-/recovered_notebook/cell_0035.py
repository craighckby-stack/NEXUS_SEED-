import os
import shutil
import argparse
import typing as t
from datetime import datetime
from docx import Document


class DistributionBuilder:
    def __init__(self, root_dir="EMG-AI-OS", package_name="EMG-AI-OS-Package"):
        self.OS_DIR = root_dir
        self.PACKAGE_DIR = package_name
        self.SCRIPTS_DATA = self._define_scripts()

    def _define_scripts(self):
        """Defines all necessary files, their content, and execution flags."""
        return {
            "bootstrap.sh": {
                "executable": True,
                "content": "#!/bin/bash\nset -e\napt update && apt install -y python3 python3-pip nftables git\necho 'EMG AI OS Bootstrap Complete.'",
            },
            "install_emg_ai.sh": {
                "executable": True,
                "content": "#!/bin/bash\nset -e\nmkdir -p /opt/emg_ai\ncp -r {os_dir}/* /opt/emg_ai/\necho 'EMG AI OS Installed Successfully.'".format(os_dir=self.OS_DIR),
            },
            "firewall.nft": {
                "executable": False,
                "content": "table inet emg_ai_firewall {\n    chain input {\n        type filter hook input priority 0;\n        policy drop;\n    }\n}",
            },
            "emg_ai.service": {
                "executable": False,
                "content": "[Unit]\nDescription=EMG AI OS Core Service\nAfter=network.target\n[Service]\nExecStart=/usr/bin/python3 /opt/emg_ai/core_service.py\nRestart=always\n[Install]\nWantedBy=multi-user.target",
            },
            "config.yaml": {
                "executable": False,
                "content": "firewall:\n  enabled: true\n  rules_file: /etc/emg_ai/firewall.nft\nkernel:\n  mode: Sovereign",
            },
            "core_service.py": {
                "executable": False,
                "content": "import time\nimport sys\nfrom datetime import datetime\nwhile True:\n    sys.stdout.write(f'[{datetime.now().isoformat()}] EMG AI OS Core Operational.\n')\n    sys.stdout.flush()\n    time.sleep(10)",
            },
        }

    def create_script_directory(self):
        print(f"[{datetime.now().strftime('%H:%M:%S')}] Preparing root directory: {self.OS_DIR}")
        os.makedirs(self.OS_DIR, exist_ok=True)

    def create_scripts(self):
        self.create_script_directory()
        print(f"[{datetime.now().strftime('%H:%M:%S')}] Writing distribution files...")
        
        for filename, data in self.SCRIPTS_DATA.items():
            filepath = os.path.join(self.OS_DIR, filename)
            content = data["content"]
            
            with open(filepath, "w") as f:
                f.write(content)
            
            if data["executable"]:
                # Ensure executable permissions (0o755)
                os.chmod(filepath, 0o755) 
                print(f"  > Wrote executable script: {filename}")
            else:
                print(f"  > Wrote data file: {filename}")

    def create_package(self):
        print(f"[{datetime.now().strftime('%H:%M:%S')}] Creating compressed package: {self.PACKAGE_DIR}.tar.gz")
        try:
            # Use gztar for gzip compression
            shutil.make_archive(self.PACKAGE_DIR, "gztar", self.OS_DIR)
        except Exception as e:
            print(f"Error creating package: {e}")

    def create_documentation(self):
        doc_path = os.path.join(self.OS_DIR, "EMG-AI-OS-Documentation.docx")
        print(f"[{datetime.now().strftime('%H:%M:%S')}] Generating documentation: {os.path.basename(doc_path)}")
        
        doc = Document()
        doc.add_heading('EMG AI OS Test Version (v94.1 Generated) - Documentation', 0)
        doc.add_paragraph(f'Build Date: {datetime.now().isoformat()}')
        doc.add_heading('Installation Steps:', level=1)
        doc.add_paragraph('1. Extract the package: `tar -xzf {package_name}.tar.gz`'.format(package_name=self.PACKAGE_DIR))
        doc.add_paragraph('2. Run bootstrap, then install components using the provided scripts.')
        doc.save(doc_path)

    def create_source_stubs(self, count=5):
        # Simulation of complex subsystem modules instead of 1000 generic files.
        print(f"[{datetime.now().strftime('%H:%M:%S')}] Creating {count} system module stubs.")
        
        module_names = ["kernel_interface", "neural_processor", "memory_manager", "security_layer", "telemetry_interface"]
        
        for i, module in enumerate(module_names[:count]):
            filename = f"module_{module}.py"
            content = f"# EMG AI OS Subsystem: {module.replace('_', ' ').title()}\n\nclass {module.title().replace('_', '')}():\n    \"\"\"Placeholder for v94.1 core logic.\"\"\"\n    def __init__(self):\n        pass"
            with open(os.path.join(self.OS_DIR, filename), "w") as f:
                f.write(content)

    def build(self):
        print("--- EMG AI OS Distribution Builder v94.1 ---")
        
        self.create_scripts()
        self.create_source_stubs(count=5)
        self.create_documentation()
        self.create_package()
        
        print(f"\n[SUCCESS] Package created: {self.PACKAGE_DIR}.tar.gz")


def main():
    # Initialize and execute build orchestration
    builder = DistributionBuilder()
    builder.build()


if __name__ == "__main__":
    # Minimal argument parsing preparation, allowing easy expansion later
    parser = argparse.ArgumentParser(description="EMG AI OS Distribution Builder")
    args = parser.parse_args() # Currently unused, placeholder for future control knobs
    main()