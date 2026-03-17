import logging
from typing import Dict, Any
from docx import Document

def generate_word_report(all_data: Dict[str, Any], output_file: str = "council_data_report.docx") -> None:
    """
    Generate a Word report from the provided council data.

    Args:
    - all_data (Dict[str, Any]): A dictionary containing council data.
    - output_file (str): The output file name. Defaults to "council_data_report.docx".
    """
    try:
        document = Document()
        document.add_heading('Council Data Analysis Report', 0)
        
        for council_name, council_data in all_data.items():
            document.add_heading(f'Council: {council_name}', 1)
            
            if 'snapshots' in council_data:
                for timestamp, pdf_info_list in council_data['snapshots'].items():
                    document.add_heading(f'Timestamp: {timestamp}', 2)
                    
                    for pdf_info in pdf_info_list:
                        document.add_heading(f"PDF: {pdf_info.get('pdf_url', 'N/A')}", 3)
                        
                        if 'checksum' in pdf_info:
                            document.add_paragraph(f"Checksum: {pdf_info['checksum']}")
                        else:
                            document.add_paragraph("Checksum not calculated.")
                        
                        if 'financial_data' in pdf_info and pdf_info['financial_data']:
                            document.add_paragraph("Financial Data:")
                            
                            for item in pdf_info['financial_data']:
                                document.add_paragraph(f"Amount: {item.get('amount', 'N/A')}")
                        else:
                            document.add_paragraph("No financial data found.")
                        
                        if 'anomalies' in pdf_info:
                            # Add anomaly information to the report
                            document.add_paragraph("Anomalies:")
                            for anomaly in pdf_info['anomalies']:
                                document.add_paragraph(f"Anomaly: {anomaly}")
        
        document.save(output_file)
        
    except Exception as e:
        logging.error(f"Failed to generate Word report: {e}")

def check_data_consistency(all_financial_data: Dict[str, Any]) -> list:
    """
    Check the consistency of the financial data.

    Args:
    - all_financial_data (Dict[str, Any]): A dictionary containing financial data for each council.

    Returns:
    - list: A list of inconsistencies found in the data.
    """
    inconsistencies = []
    
    for council_name, financial_data in all_financial_data.items():
        # Check for inconsistencies in the financial data
        for i in range(len(financial_data) - 1):
            if financial_data[i].get('amount', 0) > financial_data[i + 1].get('amount', 0):
                inconsistencies.append(f"Inconsistency in {council_name} financial data: {financial_data[i]} and {financial_data[i + 1]}")
    
    return inconsistencies

def write_to_audit_trail(event_type: str, event_description: str) -> None:
    """
    Write an event to the audit trail.

    Args:
    - event_type (str): The type of event.
    - event_description (str): A description of the event.
    """
    try:
        with open("audit_trail.log", "a") as f:
            f.write(f"{event_type}: {event_description}\n")
    
    except Exception as e:
        logging.error(f"Failed to write to audit trail: {e}")

def main() -> None:
    all_data = {}
    council_name = "Example Council"
    timestamp = "2022-01-01"
    pdf_info = {
        'pdf_url': 'https://example.com/pdf',
        'checksum': 'abc123',
        'financial_data': [
            {'amount': 100.0},
            {'amount': 200.0}
        ],
        'anomalies': ['Anomaly 1', 'Anomaly 2']
    }
    
    if council_name not in all_data:
        all_data[council_name] = {'snapshots': {}}
    
    if timestamp not in all_data[council_name]['snapshots']:
        all_data[council_name]['snapshots'][timestamp] = []
    
    all_data[council_name]['snapshots'][timestamp].append(pdf_info)
    
    all_financial_data = {council_name: [] for council_name in all_data}
    
    for council_name, council_data in all_data.items():
        for timestamp, pdf_info_list in council_data.get('snapshots', {}).items():
            for pdf_info in pdf_info_list:
                all_financial_data[council_name].append(pdf_info)
    
    inconsistencies = check_data_consistency(all_financial_data)
    
    if inconsistencies:
        logging.warning("[CONSISTENCY] Data inconsistencies found:")
        
        for inconsistency in inconsistencies:
            logging.warning(f"[CONSISTENCY] {inconsistency}")
            write_to_audit_trail("Data Inconsistency", inconsistency)
    
    generate_word_report(all_data)
    logging.info("Analysis complete.")

if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    main()