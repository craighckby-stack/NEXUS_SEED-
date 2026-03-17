import os
import yaml
import shutil
from docx import Document
import google.colab.files as files

# Define constants
MIA_OS_DIR = "MIA-OS"
PACKAGE_NAME = "MIA-OS-Package"
DOCUMENTATION_NAME = "MIA-OS-Documentation.docx"
PACKAGE_FILE_NAME = f"{PACKAGE_NAME}.tar.gz"

# Define file structure
os.makedirs(MIA_OS_DIR, exist_ok=True)

# Define scripts and configurations (Improved structure: includes install script and service unit)
scripts = {
    # 1. Configuration file (YAML)
    "firewall.yaml": {
        "firewall": {
            "enabled": True,
            "rules_file": "/etc/mia/firewall.nft",
            "version": "1.0.0-alpha" 
        }
    },
    # 2. Core Python service
    "core_service.py": """import time\nimport datetime\n\nwhile True:\n    print(f'[{datetime.datetime.now()}] MIA OS heartbeat (v1.0.0)')\n    time.sleep(60)""",
    # 3. Initial system bootstrap
    "bootstrap.sh": """#!/bin/bash\nset -e\necho \"Running initial bootstrap...\"\n# Ensure environment is available for installation\napt update -qq && apt install -y python3 python3-yaml systemd\n""",
    
    # 4. Main installation script (Previously missing, crucial for flow)
    "install_mia.sh": """#!/bin/bash\nset -e\necho \"Starting MIA OS installation...\"\n\nTARGET_DIR=\"/opt/mia\"\nCONFIG_DIR=\"/etc/mia\"\nSYSTEMD_DIR=\"/etc/systemd/system\"\n\n# 1. Setup standard directories\nmkdir -p $TARGET_DIR\nmkdir -p $CONFIG_DIR\n\n# 2. Copy core files from the unpacked package (current directory)\ncp core_service.py $TARGET_DIR/\ncp firewall.yaml $CONFIG_DIR/\nchmod +x $TARGET_DIR/core_service.py\n\n# 3. Setup systemd service\ncp mia.service $SYSTEMD_DIR/\nsystemctl daemon-reload\nsystemctl enable mia.service\n\necho \"MIA OS installation complete.\"\n""",
    
    # 5. Systemd unit file (Required for `systemctl start mia.service`)
    "mia.service": """[Unit]\nDescription=MIA OS Core Service\nAfter=network.target\n\n[Service]\nType=simple\nExecStart=/usr/bin/python3 /opt/mia/core_service.py\nRestart=always\nUser=root \nGroup=root\n\n[Install]\nWantedBy=multi-user.target\n"""
}

# Define installation steps (now logically complete)
installation_steps = {
    "Extract the package": f"`tar -xzf {PACKAGE_FILE_NAME}`",
    "Run the bootstrap script": "bash bootstrap.sh",
    "Install MIA OS": "bash install_mia.sh",
    "Start the service": "systemctl start mia.service"
}

# Create and save files inside the directory
for filename, content in scripts.items():
    path = f"{MIA_OS_DIR}/{filename}
    if isinstance(content, dict):
        with open(path, "w") as f:
            yaml.dump(content, f)
    else:
        with open(path, "w") as f:
            f.write(content)

# Package files
shutil.make_archive(PACKAGE_NAME, 'gztar', MIA_OS_DIR)

# Generate documentation
doc = Document()
doc.add_heading('MIA OS Test Version - Documentation', 0)
doc.add_paragraph('This document contains installation steps and details for MIA OS (v1.0.0-alpha).')
doc.add_heading('Installation Steps:', level=1)
for step, instructions in installation_steps.items():
    doc.add_paragraph(step)
    doc.add_paragraph(instructions)
doc.save(DOCUMENTATION_NAME)

# Output files for download (Colab specific)
files.download(PACKAGE_FILE_NAME)
files.download(DOCUMENTATION_NAME)