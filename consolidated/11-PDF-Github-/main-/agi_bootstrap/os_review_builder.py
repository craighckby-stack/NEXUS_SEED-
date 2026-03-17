import os
import datetime
from docx import Document

class OSReviewBuilder:
    """ Builds OS review packages with analysis markers """

    def __init__(self, os_files_dir):
        self.os_files_dir = os_files_dir
        self.document = Document()

    def create_review_package(self) -> str:
        """Creates a timestamped directory with OS files and a review guide"
        review_dir = f"OS_Review_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}"
        os.makedirs(review_dir, exist_ok=True)

        for file in os.listdir(self.os_files_dir):
            src = os.path.join(self.os_files_dir, file)
            dest = os.path.join(review_dir, file)
            self._add_review_markers(src, dest)

        self._generate_analysis_guide(review_dir)
        self.save_review_guide(review_dir)
        return review_dir

    def _add_review_markers(self, src: str, dest: str) -> None:
        """Adds review comments to source files"
        with open(src, 'r') as f_in, open(dest, 'w') as f_out:
            content = f_in.read()
            marked_content = f"/* REVIEW REQUIRED: {datetime.datetime.now()} */\n{content}"
            f_out.write(marked_content)

    def _generate_analysis_guide(self, review_dir: str) -> None:
        """Creates a structured review document"
        self.document.add_heading('OS Structure Analysis Report', 0)
        self.document.add_paragraph(f"Generated: {datetime.datetime.now()}"

        review_table = self.document.add_table(rows=1, cols=4)
        hdr_cells = review_table.rows[0].cells
        hdr_cells[0].text = 'Component'
        hdr_cells[1].text = 'Review Status'
        hdr_cells[2].text = 'Security Rating'
        hdr_cells[3].text = 'Notes'

        for file in os.listdir(review_dir):
            row_cells = review_table.add_row().cells
            row_cells[0].text = file
            row_cells[1].text = "Pending"
            row_cells[2].text = "Unrated"
            row_cells[3].text = "Needs initial analysis"

    def save_review_guide(self, review_dir: str) -> None:
        """Saves the review guide to a docx file"
        self.document.save(os.path.join(review_dir, 'Review_Guide.docx'))"""}
class SecurityAlertSystem:
    """Non-destructive alert system with activity logging"

class SecurityAlertSystem:
    def __init__(self):
        self.log_file = "security_audit.log"

    def log_security_activity(self, message: str) -> None:
        """Logs security-related activity"
        with open(self.log_file, 'a') as f:
            f.write(message + \n")
class AlertPrinter:
    """Prints alert messages to the console"

class AlertPrinter:
    def __init__(self):
        pass

    def print_security_alert(self, message: str) -> None:
        """Prints a security alert message"
        print(f"   \u2705  {message}  \u2705 ")
