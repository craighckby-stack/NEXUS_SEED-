import time
from urllib.parse import urlsplit
from bs4 import BeautifulSoup
import requests
from docx import Document
import re

# --- Configuration & State Management ---
SCRAPING_CONFIG = {
    'TIMEOUT': 10, 
    'DELAY_SECONDS': 1, 
    'HEADERS': {
        'User-Agent': 'SovereignAGI/v94.1 CodebaseEvolution (contact@example.com)'
    }
}
SESSION = requests.Session()
SESSION.headers.update(SCRAPING_CONFIG['HEADERS'])

# --- Utility and Log Analysis Functions ---

def extract_urls(log_data):
    """
    Extracts URLs from log data by identifying strings starting with 'http'.
    More robust than simply grabbing the last token.
    """
    urls = []
    for line in log_data:
        parts = line.split()
        for part in parts:
            if part.startswith(("http://", "https://")):
                urls.append(part.strip())
                break
    return urls

def find_threat_indicators(log_data):
    """Finds threat indicators using expanded keywords in log data, case insensitive."""
    threat_indicators = []
    KEYWORDS = ["threat", "malware", "exploit", "unauthorized", "inject", "403 forbidden"]
    for line in log_data:
        line_lower = line.lower()
        if any(kw in line_lower for kw in KEYWORDS):
            threat_indicators.append(line)
    return threat_indicators

# --- Network and Scraper Functions ---

def extract_urls_from_sitemap(sitemap_url):
    """Extracts URLs from a sitemap (XML format), using session and configured timeout."""
    try:
        response = SESSION.get(sitemap_url, timeout=SCRAPING_CONFIG['TIMEOUT'])
        response.raise_for_status()
        
        soup = BeautifulSoup(response.content, 'xml')
        urls = [loc.text for loc in soup.find_all('loc')]
        
        return urls
    except requests.RequestException as e:
        print(f"Error requesting sitemap {sitemap_url}: {e}")
        return []
    except Exception as e:
        print(f"Error parsing sitemap {sitemap_url}: {e}")
        return []

def analyze_robots_txt_directives(base_url):
    """Analyzes robots.txt directives for a URL, returning raw directives if found."""
    robots_url = f"{base_url.rstrip('/')}/robots.txt"
    try:
        response = SESSION.get(robots_url, timeout=SCRAPING_CONFIG['TIMEOUT'])
        response.raise_for_status()
        
        directives = response.text.strip().split('\n')
        print(f"Analyzing robots.txt for {base_url}: Found {len(directives)} lines.")
        return directives
        
    except requests.exceptions.HTTPError as e:
        if e.response.status_code == 404:
            print(f"robots.txt not found for {base_url}.")
        else:
            print(f"HTTP Error analyzing robots.txt for {base_url}: {e}")
        return []
    except requests.RequestException as e:
        print(f"Network Error analyzing robots.txt for {base_url}: {e}")
        return []

def scrape_page_for_info(url, document):
    """Scrapes a page for structured information (Title, H1, Summary Paragraphs) and reports to the document."""
    try:
        response = SESSION.get(url, timeout=SCRAPING_CONFIG['TIMEOUT'])
        response.raise_for_status() 

        soup = BeautifulSoup(response.content, 'html.parser')
        
        title_tag = soup.find('title')
        title = title_tag.text.strip() if title_tag else "No Title Found"

        document.add_heading(f"URL Analysis: {url}", level=2)
        document.add_paragraph(f"Title: {title}")
        document.add_paragraph(f"Status Code: {response.status_code}")

        # Extracting key content summary
        content_summary = []
        h1 = soup.find('h1')
        if h1:
            content_summary.append(f"H1: {h1.text.strip()}")

        paragraphs = soup.find_all('p', limit=3)
        for i, p in enumerate(paragraphs):
            text = re.sub(r'\s+', ' ', p.text).strip()
            if text:
                content_summary.append(f"P{i+1}: {text[:200]}...")

        if content_summary:
            document.add_paragraph("Extracted Content:")
            for item in content_summary:
                document.add_paragraph(item)
        else:
            document.add_paragraph("No significant text content extracted.")

    except requests.exceptions.RequestException as e:
        error_msg = f"Network or HTTP error during scraping {url}: {e}"
        document.add_paragraph(f"Error (HTTP/Network): {error_msg}")
        print(error_msg)
    except Exception as e:
        error_msg = f"Parsing error occurred scraping {url}: {e}"
        document.add_paragraph(f"Error (Parsing): {error_msg}")
        print(error_msg)

def process_urls_and_scrape(urls_to_scrape, document):
    """Processes a list of URLs, managing rate limits and delegating scraping."""
    document.add_heading(f"Detailed Web Scraping Results ({len(urls_to_scrape)} URLs)", level=1)
    
    for i, url in enumerate(urls_to_scrape):
        print(f"Scraping URL {i+1}/{len(urls_to_scrape)}: {url}")
             
        scrape_page_for_info(url, document)
        
        if i < len(urls_to_scrape) - 1:
            time.sleep(SCRAPING_CONFIG['DELAY_SECONDS'])

# --- Orchestration ---

def run_analysis(initial_log_data, sitemap_list):
    """Orchestrates the entire data extraction and scraping process."""
    
    extracted_urls = extract_urls(initial_log_data)
    threat_indicators = find_threat_indicators(initial_log_data)
    
    document = Document()
    document.add_heading("Comprehensive Reconnaissance Report", level=0)
    
    document.add_paragraph(f"Total extracted URLs from logs: {len(extracted_urls)}")
    if threat_indicators:
        document.add_paragraph(f"Found {len(threat_indicators)} potential threat indicators.")

    # Sitemap Harvesting
    all_harvested_urls = []
    for sitemap_url in sitemap_list:
        all_harvested_urls.extend(extract_urls_from_sitemap(sitemap_url))

    document.add_paragraph(f"\nTotal URLs harvested from {len(sitemap_list)} sitemaps: {len(all_harvested_urls)}")

    # Robots.txt Analysis
    base_urls = set()
    for url in extracted_urls:
        parsed = urlsplit(url)
        if parsed.netloc:
            scheme = parsed.scheme if parsed.scheme else 'https'
            base_urls.add(f"{scheme}://{parsed.netloc}")
            
    document.add_heading(f"Robots.txt Directives for {len(base_urls)} Domains", level=1)
    for url in base_urls:
        analyze_robots_txt_directives(url)

    # Detailed Scraping
    process_urls_and_scrape(extracted_urls, document)
    
    output_filename = 'sovereign_v94_scraping_results.docx'
    document.save(output_filename)
    print(f"\nAnalysis complete. Results saved to {output_filename}")


if __name__ == "__main__":
    log_data = [
        "192.168.1.1 - [10/Oct/2000:13:55:36] GET /index.html HTTP/1.1 200 2326 http://example.com/referrer.html", 
        "Suspicious activity detected: threat indicator payload found.",
        "Request failed for resource: https://example.org/api/data?q=123",
        "Just a normal line without a URL."
    ]
    
    sitemap_urls = [
        "https://sbi.co.in/webfiles/uploads/files_2122/sitemap.xml",
        "https://sbi.co.in/webfiles/uploads/files_2122/sitemap-cards.xml",
        "https://sbi.co.in/webfiles/uploads/files_2122/sitemap-investments-deposits.xml",
        "https://sbi.co.in/webfiles/uploads/files_2122/sitemap-loans.xml",
        "https://sbi.co.in/webfiles/uploads/files_2122/sitemap-main.xml",
        "https://sbi.co.in/webfiles/uploads/files_2122/sitemap-saving.xml",
        "https://canarabank.com/sitemap.xml",
        "https://canarabank.com/images_sitemap.xml",
        "https://bankofindia.co.in/sitemap.xml"
    ]

    run_analysis(log_data, sitemap_urls)