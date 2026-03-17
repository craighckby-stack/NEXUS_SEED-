import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse, urlunparse
from nmap import PortScanner, PortScannerError
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any, Optional, Callable
from concurrent.futures import ThreadPoolExecutor, as_completed

# --- Configuration Constants (Centralized and Typed) ---
class ReconConfig:
    DEFAULT_HTTP_TIMEOUT: int = 5
    DEFAULT_CONTENT_TIMEOUT: int = 15 # Slightly longer for large pages
    MAX_WORKERS: int = 5 # For parallel fetching
    NMAP_PORTS: str = '80,443,8080'
    NMAP_ARGS: str = '-sV -T4 --script=http-title'

CONFIG = ReconConfig()

class ReconnaissanceAgent:
    """Manages all reconnaissance tasks for a given target URL, using concurrency for speed."""

    def __init__(self, target_url: str):
        # Ensure the URL has a scheme for reliable parsing and fetching
        if not urlparse(target_url).scheme:
            target_url = f"http://{target_url}"
            
        self.target_url = target_url
        parsed = urlparse(target_url)
        # Robustly extract the host, falling back to path if netloc is empty (e.g. raw IP)
        self.target_host = parsed.netloc if parsed.netloc else parsed.path

    def analyze_security_headers(self, headers: Dict[str, str]) -> Dict[str, Any]:
        """Checks for critical security headers and provides a simple report."""
        security_checks = {}
        
        # Standard headers are already lowercased in fetch_http_headers
        
        security_checks['HSTS'] = {
            'present': 'strict-transport-security' in headers,
            'value': headers.get('strict-transport-security', 'N/A')
        }
        security_checks['X-Frame-Options'] = {
            'present': 'x-frame-options' in headers,
            'value': headers.get('x-frame-options', 'N/A')
        }
        security_checks['Content-Security-Policy'] = {
            'present': 'content-security-policy' in headers,
            'value': headers.get('content-security-policy', 'N/A')[:100] + ('...' if len(headers.get('content-security-policy', '')) > 100 else '')
        }
        security_checks['X-Content-Type-Options'] = {
            'present': 'x-content-type-options' in headers,
            'value': headers.get('x-content-type-options', 'N/A')
        }

        missing_count = sum(1 for check in security_checks.values() if not check['present'])
        
        return {"score": f"{4 - missing_count}/4 Secure Headers Present", "checks": security_checks}

    def fetch_robots_txt(self) -> Dict[str, Any]:
        """Fetches the robots.txt content and returns structured data."""
        results: Dict[str, Any] = {}
        # print(f"[INFO]: Fetching robots.txt for {self.target_host}")
        try:
            robots_url = urljoin(self.target_url, 'robots.txt')
            # We use a short timeout as this file should be small
            response = requests.get(robots_url, timeout=CONFIG.DEFAULT_HTTP_TIMEOUT, allow_redirects=False)
            results['status_code'] = response.status_code
            
            if response.status_code == 200:
                results['content'] = response.text
                results['size_bytes'] = len(response.content)
            elif response.status_code == 404:
                results['content'] = "File not found (404)"
            else:
                results['content'] = response.text[:200] # Preview for non-200, non-404 status
                
            return results
        except requests.exceptions.RequestException as e:
            return {"error": f"Error fetching robots.txt: {e}", "status_code": 0}

    def fetch_http_headers(self) -> Dict[str, str]:
        """Fetches and returns the HTTP headers dictionary using a context manager."""
        # print(f"[INFO]: Fetching headers for {self.target_host}")
        
        try:
            # Using context manager for safe resource handling
            # Using HEAD request might be faster but GET ensures full header set often
            with requests.get(self.target_url, timeout=CONFIG.DEFAULT_HTTP_TIMEOUT) as response:
                response.raise_for_status()
                # Standardize headers to lowercase for easier security analysis/reporting
                return {k.lower(): v for k, v in response.headers.items()}
                
        except requests.exceptions.RequestException as e:
            # Ensure returning a dict compatible with later post-processing
            return {"error": f"Error fetching HTTP headers: {e}"}

    def fetch_webpage_content(self) -> Dict[str, Any]:
        """Fetches detailed content metadata (links, form count, SEO metadata)."""
        # print(f"[INFO]: Analyzing page content for {self.target_host}")
        try:
            response = requests.get(self.target_url, timeout=CONFIG.DEFAULT_CONTENT_TIMEOUT)
            response.raise_for_status()
            soup = BeautifulSoup(response.text, 'html.parser')

            # Filter links: keep external, absolute internal, or relative internal links
            links = [
                link['href'] for link in soup.find_all('a', href=True)
                if link['href'].startswith(('http', '/', './', '../')) # Simplified link filter
            ]
            forms = soup.find_all('form')
            
            # Metadata Extraction
            metadata = {}
            for tag in soup.find_all('meta'):
                if tag.get('name') and tag.get('content'):
                    metadata[tag['name'].lower()] = tag['content'].strip()
                elif tag.get('property') and tag.get('content'): # Handle OpenGraph and other proprietary standards
                    metadata[tag['property'].lower()] = tag['content'].strip()
            
            # Get Canonical URL if present
            canonical = soup.find('link', {'rel': 'canonical'})
            canonical_url = canonical['href'] if canonical and canonical.get('href') else None
            
            return {
                "title": soup.title.string.strip() if soup.title and soup.title.string else "No Title Found",
                "canonical_url": canonical_url,
                "total_links": len(links),
                "form_count": len(forms),
                "metadata": metadata,
                "sample_links": links[:10],
            }
        except requests.exceptions.RequestException as e:
            return {"error": f"Error fetching webpage content: {e}"}
        except Exception as e:
             return {"error": f"Error parsing content: {e}"}

    def run_nmap_scan(self) -> Dict[str, Any]:
        """Performs a basic Nmap scan and returns structured results."""
        scan_results: Dict[str, Any] = {}
        
        if not self.target_host:
            return {"error": "Invalid target host provided for Nmap scan."}
            
        # print(f"[INFO]: Starting Nmap Scan on {self.target_host}")
        
        try:
            nm = PortScanner()
            
            # Use CONFIG
            nm.scan(self.target_host, arguments=f'{CONFIG.NMAP_ARGS} -p {CONFIG.NMAP_PORTS}')  
            
            for host in nm.all_hosts():
                host_info = {
                    "status": nm[host].state(),
                    "protocols": {}
                }
                
                if not nm[host]:
                    continue
                    
                for proto in nm[host].all_protocols():
                    proto_info = {}
                    for port in sorted(nm[host][proto].keys()):
                        port_data = nm[host][proto][port]
                        proto_info[port] = {
                            "state": port_data.get('state', 'unknown'),
                            "name": port_data.get('name', 'N/A'),
                            "product": port_data.get('product', 'N/A'),
                            "version": port_data.get('version', 'N/A'),
                            "extra_info": port_data.get('extrainfo', '') # Capture richer detail
                        }
                    host_info['protocols'][proto] = proto_info
                scan_results[host] = host_info

        except PortScannerError as e:
            scan_results["error"] = f"Nmap scanning configuration error (Is Nmap installed?): {e}"
        except Exception as e:
            scan_results["error"] = f"An unexpected error occurred during the Nmap scan: {e}"
            
        return scan_results

    def execute_recon(self) -> Dict[str, Any]:
        """Runs all configured reconnaissance steps concurrently and aggregates results."""
        recon_data = {
            "target": self.target_url,
            "host": self.target_host
        }

        if not self.target_host:
            return {"error": "Target URL is malformed or missing host identifier.", "target": self.target_url}
        
        print(f"--- Starting Concurrent Reconnaissance on {self.target_host} ---")
        
        tasks: Dict[str, Callable[[], Dict[str, Any]]] = {
            'robots': self.fetch_robots_txt,
            'headers': self.fetch_http_headers,
            'content': self.fetch_webpage_content,
            'nmap': self.run_nmap_scan,
        }

        # Use ThreadPoolExecutor for I/O bound tasks (HTTP fetches + Nmap)
        with ThreadPoolExecutor(max_workers=CONFIG.MAX_WORKERS) as executor:
            future_to_task = {executor.submit(func): name for name, func in tasks.items()}
            
            for future in as_completed(future_to_task):
                task_name = future_to_task[future]
                try:
                    print(f"[INFO]: Finished {task_name} check.")
                    recon_data[task_name] = future.result()
                except Exception as exc:
                    print(f"[ERROR]: Task {task_name} generated an unhandled exception: {exc}")
                    recon_data[task_name] = {"error": f"Runtime error: {exc}"}

        # Post-processing / Synchronization Step: Analyze security headers
        if 'headers' in recon_data and 'error' not in recon_data['headers']:
            print("[INFO]: Analyzing Security Headers...")
            recon_data['security_analysis'] = self.analyze_security_headers(recon_data['headers'])
        else:
             recon_data['security_analysis'] = {"error": "Skipped due to header retrieval failure."}
        
        print(f"--- Reconnaissance Complete ---")
        return recon_data

def generate_report(data: Dict[str, Any], filename: str = "recon_report.docx") -> None:
    """Creates a DOCX report from the collected reconnaissance data."""
    doc = Document()
    doc.add_heading('Sovereign AGI Target Reconnaissance Report', 0)
    
    # Target Info
    doc.add_paragraph(f'Target URL: {data["target"]}')
    doc.add_paragraph(f'Target Host: {data["host"]}')

    # 1. Web Content Summary & Metadata
    doc.add_heading('1. Web Content Summary', 2)
    web_summary = data.get('content', {})
    if 'error' in web_summary:
         doc.add_paragraph(f"Status: Failed to analyze content. Error: {web_summary['error']}")
    else:
        doc.add_paragraph(f'Page Title: {web_summary.get("title")}')
        doc.add_paragraph(f'Canonical URL: {web_summary.get("canonical_url", "N/A")}')
        doc.add_paragraph(f'Total Links Found: {web_summary.get("total_links")}')
        doc.add_paragraph(f'Forms Found: {web_summary.get("form_count")}')
        
        # Metadata Table
        if web_summary.get('metadata'):
            doc.add_heading('1.1 Page Metadata', 3)
            p = doc.add_paragraph()
            for key, value in web_summary['metadata'].items():
                p.add_run(f'{key}: ').bold = True
                p.add_run(f'{value[:100]}...
') # Truncate long descriptions

    
    # 2. HTTP Headers & Security Analysis
    doc.add_heading('2. HTTP Headers & Security Analysis', 2)
    headers_data = data.get('headers', {})
    security_data = data.get('security_analysis', {})
    
    if 'error' in headers_data:
         doc.add_paragraph(f"Status: Failed to retrieve headers. Error: {headers_data['error']}")
    else:
        # 2.1 Security Summary
        doc.add_heading('2.1 Security Header Score', 3)
        doc.add_paragraph(f"Score: {security_data.get('score', 'N/A')}").bold = True
        
        sec_table = doc.add_table(rows=1, cols=3)
        sec_table.style = 'Light Grid'
        hdr_cells = sec_table.rows[0].cells
        hdr_cells[0].text = 'Header'
        hdr_cells[1].text = 'Present?'
        hdr_cells[2].text = 'Value (Preview)'

        if 'checks' in security_data:
            for key, check in security_data['checks'].items():
                row_cells = sec_table.add_row().cells
                row_cells[0].text = key
                row_cells[1].text = 'Yes' if check['present'] else 'NO'
                row_cells[2].text = check['value']

        # 2.2 Raw Headers
        doc.add_heading('2.2 Full HTTP Headers', 3)
        # Create a simple table for headers
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Header Key'
        hdr_cells[1].text = 'Value'
        
        for key, value in headers_data.items():
            # Skip long header values like Content-Security-Policy that are in the analysis section
            if key not in ['strict-transport-security', 'x-frame-options', 'content-security-policy', 'x-content-type-options']:
                row_cells = table.add_row().cells
                row_cells[0].text = key
                row_cells[1].text = value

    
    # 3. Robots.txt
    doc.add_heading('3. Robots.txt Status', 2)
    robots = data.get('robots', {})
    if 'error' in robots:
         doc.add_paragraph(f"Status: Failed to retrieve. Error: {robots['error']}")
    else:
        doc.add_paragraph(f"Status Code: {robots.get('status_code', 'N/A')}")
        if robots.get('content') and 'File not found' not in robots['content']:
            doc.add_paragraph(f"Content Preview (500 chars):
{robots['content'][:500]}...")
        else:
            doc.add_paragraph(robots.get('content', 'No content retrieved.'))

    # 4. Nmap Scan Results
    doc.add_heading('4. Nmap Scan Results', 2)
    nmap_results = data.get('nmap', {})
    if 'error' in nmap_results:
        doc.add_paragraph(f"Scan Error: {nmap_results['error']}")
    else:
        if not nmap_results:
            doc.add_paragraph("Nmap scan completed, but no live hosts were reported on specified ports.")
        for host, host_info in nmap_results.items():
            doc.add_heading(f'Host: {host} (Status: {host_info["status"].upper()})', 3)
            
            for proto, ports in host_info.get('protocols', {}).items():
                doc.add_paragraph(f'Protocol: {proto.upper()}')
                
                table = doc.add_table(rows=1, cols=5) # Added Extra Info column
                table.style = 'Light Shading Accent 1'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Port'
                hdr_cells[1].text = 'State'
                hdr_cells[2].text = 'Service/Product'
                hdr_cells[3].text = 'Version'
                hdr_cells[4].text = 'Extra Info'
                
                for port, info in ports.items():
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(port)
                    row_cells[1].text = info['state']
                    row_cells[2].text = info.get('product') or info.get('name') or 'N/A'
                    row_cells[3].text = info['version']
                    row_cells[4].text = info.get('extra_info', '')[:50]

    doc.save(filename)
    print(f"\n[REPORTING]: Report saved successfully to {filename}")

def main():
    target_url = "http://example.com"
    
    agent = ReconnaissanceAgent(target_url)
    
    if not agent.target_host:
        print(f"Error: Invalid target identification: {target_url}")
        return

    # 1. Gather Data
    recon_data = agent.execute_recon()
    
    # 2. Summary Output (For CLI/Notebook viewing)
    print("\n=================== SUMMARY ===================")
    print(f"Target URL: {recon_data['target']}")
    
    content_data = recon_data.get('content', {})
    nmap_data = recon_data.get('nmap', {})
    header_data = recon_data.get('headers', {})
    security_data = recon_data.get('security_analysis', {})

    print(f"Web Title: {content_data.get('title', content_data.get('error', 'N/A'))}")
    
    server_header = header_data.get('server', header_data.get('error', 'N/A'))
    print(f"Server Header: {server_header}")
    print(f"Security Headers Score: {security_data.get('score', 'N/A')}")

    if 'error' not in nmap_data:
         print(f"Nmap hosts found: {list(nmap_data.keys())}")
    else:
         print(f"Nmap Scan Failed: {nmap_data['error']}")

    # 3. Generate structured report
    generate_report(recon_data)

if __name__ == "__main__":
    main()