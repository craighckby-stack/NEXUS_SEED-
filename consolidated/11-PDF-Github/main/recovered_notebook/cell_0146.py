import requests
import re
import logging
from bs4 import BeautifulSoup
import time
from docx import Document
from urllib.parse import urljoin, urlparse
import concurrent.futures
from dataclasses import dataclass, field

# --- AGI v94.1 Configuration --- 
@dataclass
class SovereignConfig:
    MAX_WORKERS: int = 15 # Increased efficiency via more workers
    RATE_LIMIT_DELAY: float = 0.05 # Aggressive reduction, managed by concurrency
    DEFAULT_TIMEOUT: int = 15
    
    # Extended and specialized threat patterns
    BITCOIN_PATTERNS: dict = field(default_factory=lambda: {
        'crypto_keyword': re.compile(r'(bitcoin|btc|cryptocurrency|wallet|mining|satoshi)', re.IGNORECASE),
        'address_format': re.compile(r'(bc1|[13])[a-zA-HJ-NP-Z0-9]{25,39}', re.IGNORECASE) # Simple BTC address pattern
    })
    THREAT_KEYWORDS: list = field(default_factory=lambda: ['threat', 'injection', 'xss', 'exploit', 'phish', 'ransom'])
    LOG_LEVEL: int = logging.INFO

CONFIG = SovereignConfig()

# --- Logging Setup ---
logging.basicConfig(level=CONFIG.LOG_LEVEL, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Regular expression for robust URL extraction (improves on basic splitting)
URL_REGEX = re.compile(r'https?://(?:[-\w.]|(?:%[0-9a-fA-F]{2}))+')

def extract_urls(log_data):
    """Extracts potential URLs from iterable data using a robust RegEx approach.
    """
    urls = set()
    for line in log_data:
        found = URL_REGEX.findall(line)
        for url in found:
            urls.add(url.strip().rstrip('.')) # Clean up trailing dots/spaces
    logger.info(f"Successfully extracted {len(urls)} potential URLs from input data.")
    return list(urls)

def extract_urls_from_sitemap(sitemap_url):
    """Extracts URLs from a sitemap, handling common errors.
    Now utilizes configured timeout.
    """
    try:
        response = requests.get(sitemap_url, timeout=5) # Shorter timeout for XML fetching
        response.raise_for_status()
        soup = BeautifulSoup(response.content, 'xml')
        urls = []
        for loc in soup.find_all('loc'):
            urls.append(loc.text.strip())
        return urls
    except requests.exceptions.HTTPError as e:
        logger.warning(f"HTTP Error fetching sitemap {sitemap_url}: {e}")
        return []
    except requests.exceptions.RequestException as e:
        logger.error(f"Connection Error fetching sitemap {sitemap_url}: {e}")
        return []

def find_threat_indicators(log_data):
    """Finds simple threat indicators in log data based on global configuration.
    """
    threat_indicators = []
    lower_keywords = [k.lower() for k in CONFIG.THREAT_KEYWORDS]
    for line in log_data:
        if any(keyword in line.lower() for keyword in lower_keywords):
            threat_indicators.append(line)
    return threat_indicators

def calculate_confidence_score(text_data, patterns):
    """Calculates a confidence score based on specialized pattern matching.
    Score ranges from 0 (Clean) to 100 (High Confidence Match).
    """
    score = 0
    text_data_lower = text_data.lower()
    
    # Check for specific address formats (High Score)
    address_matches = patterns['address_format'].findall(text_data)
    if address_matches:
        score += 50 + min(len(address_matches) * 10, 30) # Max 80 just for addresses
        logger.debug(f"Address format matches: {len(address_matches)}")

    # Check for general keywords (Medium Score)
    keyword_matches = patterns['crypto_keyword'].findall(text_data_lower)
    if keyword_matches:
        score += min(len(keyword_matches) * 5, 20) # Max 20 for keywords
        logger.debug(f"Keyword matches: {len(keyword_matches)}")
        
    return min(score, 100), address_matches

def analyze_page_for_bitcoin(url):
    """Analyzes a single page for configured patterns, returning structured results with confidence.
    """
    results = {
        "url": url,
        "status": "FAILURE",
        "confidence": 0,
        "matches": [],
        "error": None
    }
    
    try:
        time.sleep(CONFIG.RATE_LIMIT_DELAY)
        response = requests.get(url, timeout=CONFIG.DEFAULT_TIMEOUT)
        response.raise_for_status()
        soup = BeautifulSoup(response.content, 'html.parser')
        page_text = soup.text
        
        score, address_matches = calculate_confidence_score(page_text, CONFIG.BITCOIN_PATTERNS)
        results['confidence'] = score

        if score >= 50:
            results['status'] = "HIGH_CONFIDENCE_MATCH"
            results['matches'] = list(set([m for m in address_matches]))
        elif score > 0:
            results['status'] = "LOW_CONFIDENCE_MATCH"
        else:
            results['status'] = "CLEAN"
            
    except requests.exceptions.HTTPError as e:
        results['error'] = f"HTTP Error {e.response.status_code}"
        logger.debug(f"HTTP Error on {url}: {e.response.status_code}")
    except requests.exceptions.RequestException as e:
        results['error'] = f"Connection/Timeout Error: {e}"
        logger.debug(f"Connection Error on {url}: {e}")
    except Exception as e:
        results['error'] = f"Unexpected error: {e}"
        logger.error(f"Unexpected error processing {url}: {e}")
        
    return results

def process_urls_and_analyze_bitcoin(urls_to_analyze):
    """Processes a list of URLs using a ThreadPoolExecutor for concurrency.
    """
    logger.info(f"Starting analysis of {len(urls_to_analyze)} URLs using {CONFIG.MAX_WORKERS} workers.")
    
    results = []
    
    with concurrent.futures.ThreadPoolExecutor(max_workers=CONFIG.MAX_WORKERS) as executor:
        future_to_url = {executor.submit(analyze_page_for_bitcoin, url): url for url in urls_to_analyze}
        
        for i, future in enumerate(concurrent.futures.as_completed(future_to_url)):
            url = future_to_url[future]
            try:
                data = future.result()
                results.append(data)
                if (i + 1) % 100 == 0 or (i + 1) == len(urls_to_analyze):
                    logger.info(f"Progress: {i + 1}/{len(urls_to_analyze)} URLs processed.")
            except Exception as exc:
                err_data = {"url": url, "status": "FAILURE", "confidence": 0, "matches": [], "error": f"Execution failed: {exc}"}
                results.append(err_data)
                logger.warning(f"Task failed for {url}: {exc}")
                
    return results

def generate_report(analysis_results, docx_document, threat_indicators):
    """Generates the docx report based on collected analysis results.
    """
    docx_document.add_heading('Sovereign AGI Bitcoin Analysis Report (v94.1)', 0)
    docx_document.add_paragraph(f"Report Generated: {time.ctime()}")
    docx_document.add_paragraph(f"Configuration Workers: {CONFIG.MAX_WORKERS}, Delay: {CONFIG.RATE_LIMIT_DELAY}s")
    
    # --- Threat Indicators ---
    docx_document.add_heading('Detected Threat Indicators (Raw Logs)', level=1)
    if threat_indicators:
        for indicator in threat_indicators:
            docx_document.add_paragraph(f"* {indicator}")
    else:
        docx_document.add_paragraph("No explicit threat keywords found in the input log data.")

    # --- Bitcoin Analysis Results ---
    docx_document.add_heading('Web Page Bitcoin Pattern Analysis', level=1)
    
    high_matches = [r for r in analysis_results if r['status'] == 'HIGH_CONFIDENCE_MATCH']
    low_matches = [r for r in analysis_results if r['status'] == 'LOW_CONFIDENCE_MATCH']
    failures = [r for r in analysis_results if r['status'] == 'FAILURE']

    docx_document.add_paragraph(f"Total URLs Analyzed: {len(analysis_results)}")
    docx_document.add_paragraph(f"High Confidence Matches (Score >= 50): {len(high_matches)}")
    docx_document.add_paragraph(f"Low Confidence Matches (Score > 0): {len(low_matches)}")
    docx_document.add_paragraph(f"Total Failures/Errors: {len(failures)}")
    
    if high_matches:
        docx_document.add_heading('Confirmed High Confidence Bitcoin/Crypto Hits', level=2)
        for match in high_matches:
            p = docx_document.add_paragraph(f"URL: {match['url']} (Confidence: {match['confidence']}%")
            p.style = 'List Bullet'
            if match['matches']:
                docx_document.add_paragraph(f"  -> Detected Addresses: {', '.join(match['matches'])}")

    if low_matches:
        docx_document.add_heading('Potential Low Confidence Hits', level=2)
        for match in low_matches:
            p = docx_document.add_paragraph(f"URL: {match['url']} (Confidence: {match['confidence']}%")
            p.style = 'List Number'
    
    if failures:
        docx_document.add_heading('Analysis Failures/Errors', level=2)
        for failure in failures:
            p = docx_document.add_paragraph(f"URL: {failure['url']}")
            p.style = 'List Number'
            docx_document.add_paragraph(f"  -> Error: {failure['error']}")


if __name__ == "__main__":
    # NOTE: Sample log_data now includes realistic mixed content and a simulated BTC address for testing the high-confidence match.
    log_data = [
        'Request 1: https://example.com/check',
        'Request 2: https://example2.com/secure',
        'Attempted injection: 127.0.0.1 - /login?user=admin; DROP TABLE users -- threat',
        'Service Announcement: Please send donations to 1A1zP1eP5QGefi2DMPTfTL5SLmv7DivfNa (Bitcoin Wallet)'
    ]
    
    logger.info("--- Sovereign AGI v94.1 Analysis Start ---")
    
    # --- 1. Preparation and Extraction ---
    document = Document()
    extracted_urls = extract_urls(log_data)
    threat_indicators = find_threat_indicators(log_data)

    # --- 2. Sitemap Analysis and URL Aggregation ---
    sitemap_urls = [
        "https://sbi.co.in/webfiles/uploads/files_2122/sitemap.xml",
        "https://bankofindia.co.in/sitemap.xml"
    ]
    all_sitemap_urls = []
    for sitemap_url in sitemap_urls:
        all_sitemap_urls.extend(extract_urls_from_sitemap(sitemap_url))

    urls_to_analyze = list(set(extracted_urls + all_sitemap_urls))
    
    # --- 3. Concurrent Analysis ---
    analysis_results = process_urls_and_analyze_bitcoin(urls_to_analyze)

    # --- 4. Report Generation ---
    generate_report(analysis_results, document, threat_indicators)
    
    filename = 'bitcoin_analysis_report_v94_1_refactored.docx'
    document.save(filename)
    logger.info(f"\nReport saved as {filename}")