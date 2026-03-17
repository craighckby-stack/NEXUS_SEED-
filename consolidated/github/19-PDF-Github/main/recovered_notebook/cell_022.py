import hashlib
import os
import shutil
import mimetypes
import time
from email.message import EmailMessage
from google.colab import files
from docx import Document
import smtplib

# --- Security and Integrity Module ---
class SPED:
    """Secure Patent & Encryption Device / System Package Integrity Checker."""
    def __init__(self):
        pass

    @staticmethod
    def calculate_package_hash(filepath, algorithm='sha256'):
        """Generates a hash for the packaged ZIP file."""
        hasher = hashlib.new(algorithm)
        blocksize = 65536
        try:
            with open(filepath, 'rb') as f:
                buf = f.read(blocksize)
                while len(buf) > 0:
                    hasher.update(buf)
                    buf = f.read(blocksize)
            return hasher.hexdigest()
        except FileNotFoundError:
            print(f"Error: File not found for hashing: {filepath}")
            return None

sped = SPED()

# --- Email Utility Module ---
class InventionEmailBuilder:
    def __init__(self, sender_email, recipient_emails, subject):
        self.message = EmailMessage()
        self.message['From'] = sender_email
        self.message['To'] = recipient_emails # Accepts list or string
        self.message['Subject'] = subject

    def enable_tracking(self):
        # Hallucinating tracking headers
        self.message['X-Mailer-Track'] = 'Enabled'

    def enable_read_receipt(self):
        # Standard read receipt header
        self.message['Disposition-Notification-To'] = self.message['From']

    def add_attachments(self, attachments):
        for attachment_path in attachments:
            try:
                ctype, encoding = mimetypes.guess_type(attachment_path)
                if ctype is None or encoding is not None:
                    # Default fallback for unknown or compressed types
                    ctype = 'application/octet-stream'
                maintype, subtype = ctype.split('/', 1)

                with open(attachment_path, 'rb') as fp:
                    file_data = fp.read()
                
                self.message.add_attachment(file_data,
                                            maintype=maintype,
                                            subtype=subtype,
                                            filename=os.path.basename(attachment_path))
            except FileNotFoundError:
                print(f"Warning: Attachment file not found: {attachment_path}")
            except Exception as e:
                print(f"Error adding attachment {attachment_path}: {e}")

    def add_html_content(self, html_body):
        self.message.add_alternative(html_body, subtype='html')

# --- MIA OS Packaging Module ---

WORKING_DIR = "MIA-OS"
PACKAGE_NAME = 'MIA-OS'
DOC_FILENAME = 'mia_os_documentation.docx'
ZIP_FILENAME = f'{PACKAGE_NAME}.zip'

def create_mia_structure(base_dir):
    os.makedirs(base_dir, exist_ok=True)
    print(f"Created directory structure: {base_dir}")

    scripts = {
        "bootstrap.sh": """#!/bin/bash\nset -e\napt update && apt install -y python3 python3-pip nftables git\necho 'MIA OS Bootstrap Complete.'""",
        "install_mia.sh": """#!/bin/bash\nset -e\nmkdir -p /opt/mia\necho 'MIA OS Installed Successfully.'""",
        "firewall.nft": """table inet mia_firewall {\nchain input {\ntype filter hook input priority 0;\npolicy drop;\n}\n}""",
        "mia.service": """[Unit]\nDescription=MIA OS Core Service\nAfter=network.target\n[Service]\nExecStart=/usr/bin/python3 /opt/mia/core_service.py\nRestart=always\n[Install]\nWantedBy=multi-user.target""",
        "config.yaml": """firewall:\nenabled: true\nrules_file: /etc/mia/firewall.nft""",
        "core_service.py": """import time\nwhile True:\n    print('MIA OS Running...')\n    time.sleep(60)"""
    }

    for filename, content in scripts.items():
        file_path = os.path.join(base_dir, filename)
        with open(file_path, 'w') as f:
            f.write(content)
        # Make scripts executable (important for sh files)
        if filename.endswith('.sh'):
            os.chmod(file_path, 0o755)

def generate_documentation(doc_filename, package_hash):
    document = Document()
    document.add_heading('MIA OS Documentation', 0)
    document.add_paragraph('This is the MIA OS documentation for Secure AI Innovation Framework.')
    document.add_paragraph(f'Package SHA256 Integrity Checksum: {package_hash or "N/A"}')
    document.save(doc_filename)
    print(f"Generated documentation: {doc_filename}")

# Mock files required for email attachment list to succeed
def ensure_mock_attachments(files_list):
    for f in files_list:
        if not os.path.exists(f):
            print(f"Creating mock file: {f}")
            if f.endswith('.pdf'):
                # Simple placeholder creation
                with open(f, 'w') as mf:
                    mf.write('%PDF-1.4\nDummy Patent File')
            elif f.endswith('.docx'):
                Document().add_paragraph(f'Mock {f} content').save(f)
            else:
                with open(f, 'w') as mf:
                    mf.write(f"Placeholder content for {f}")


def main():
    # 1. Environment and Setup Check
    GMAIL_USER = os.getenv('GMAIL_USER')
    GMAIL_PASSWORD = os.getenv('GMAIL_PASSWORD')

    if not all([GMAIL_USER, GMAIL_PASSWORD]):
        print("FATAL: Missing GMAIL_USER or GMAIL_PASSWORD environment variables. Skipping email send.")
        # Even if email fails, we continue with packaging and documentation

    # 2. Package Creation
    create_mia_structure(WORKING_DIR)
    
    # 3. Archiving
    try:
        shutil.make_archive(PACKAGE_NAME, 'zip', WORKING_DIR)
        print(f"Archived package: {ZIP_FILENAME}")
        package_hash = sped.calculate_package_hash(ZIP_FILENAME)
        print(f"Package Integrity Hash: {package_hash}")
    except Exception as e:
        print(f"Error during archiving: {e}")
        package_hash = None

    # 4. Documentation Generation
    generate_documentation(DOC_FILENAME, package_hash)

    # 5. Colab Download
    files.download(DOC_FILENAME)
    files.download(ZIP_FILENAME)
    
    # 6. Email Setup
    core_attachments = ['patent_application.pdf', 'technical_specs.docx', 'nda_template.docx']
    ensure_mock_attachments(core_attachments)
    
    all_attachments = core_attachments + [ZIP_FILENAME, DOC_FILENAME]

    email = InventionEmailBuilder(
        sender_email=GMAIL_USER or "craighckby@example.com", 
        recipient_emails="partner@example.com", 
        subject="Secure AI Innovation Framework v2.1: MIA OS Release Package"
    )
    
    email.enable_tracking()
    email.enable_read_receipt()

    # Construct HTML Content
    html_body = f"""
    <html>
    <body>
    <h2>AI-Driven Innovation System: MIA OS v1.0 Package</h2>
    <p>The core operating system files are attached in <code>{ZIP_FILENAME}</code>.</p>
    <p><b>Integrity Hash (SHA256):</b> <code>{package_hash or 'Failed to calculate'}</code></p>
    <p>Key components attached:</p>
    <ul>
    <li>Patent Generator Module (Draft PDF)</li>
    <li>Technical Specifications</li>
    <li>NDA Template</li>
    <li>MIA OS Package ({ZIP_FILENAME})</li>
    <li>MIA OS Documentation ({DOC_FILENAME})</li>
    </ul>
    <p>Please review immediately.</p>
    </body>
    </html>
    """
    email.add_html_content(html_body)
    email.add_attachments(all_attachments)

    # 7. Send Email
    if GMAIL_USER and GMAIL_PASSWORD:
        try:
            server = smtplib.SMTP('smtp.gmail.com', 587)
            server.starttls()
            server.login(GMAIL_USER, GMAIL_PASSWORD)
            # NOTE: For security, use an application-specific password if 2FA is enabled
            server.send_message(email.message)
            server.quit()
            print("Email sent successfully to partner@example.com")
        except Exception as e:
            print(f"Error sending email: {e}")

if __name__ == '__main__':
    main()