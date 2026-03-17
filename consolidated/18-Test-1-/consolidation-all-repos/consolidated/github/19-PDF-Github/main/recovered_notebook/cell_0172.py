import logging
import traceback
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.table import Table

# --- Configuration & Architectural Definition ---
REPORT_CONFIG = {
    'TITLE_FONT_SIZE': 24,
    'DEFAULT_TABLE_STYLE': 'LightShadingAccent1',
    'FINANCIAL_COLUMNS': [
        {'header': 'Vendor', 'width': 1.5, 'alignment': WD_ALIGN_PARAGRAPH.LEFT},
        {'header': 'Amount', 'width': 1.0, 'alignment': WD_ALIGN_PARAGRAPH.RIGHT},
        {'header': 'Currency', 'width': 0.75, 'alignment': WD_ALIGN_PARAGRAPH.LEFT},
        {'header': 'Anomaly Score', 'width': 1.25, 'alignment': WD_ALIGN_PARAGRAPH.RIGHT}
    ]
}

# Configuration
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')


def _add_report_title(document, title):
    """Adds a primary title, subtitle, and metadata (date) to the document."""
    # Primary Title
    heading = document.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if heading.runs:
        heading.runs[0].font.size = Pt(REPORT_CONFIG['TITLE_FONT_SIZE'])

    # Subtitle / Metadata
    date_p = document.add_paragraph()
    date_p.add_run(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").italic = True
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph() # Add spacer


def _add_financial_data_table(document, financial_data):
    """Adds the financial data table section with configuration-driven formatting."""
    document.add_heading('1. Financial Transactions Summary', level=2)

    if not financial_data:
        document.add_paragraph("No financial data records found for this analysis.")
        return

    config_cols = REPORT_CONFIG['FINANCIAL_COLUMNS']
    num_cols = len(config_cols)
    
    table = document.add_table(rows=1, cols=num_cols, style=REPORT_CONFIG['DEFAULT_TABLE_STYLE'])
    
    # Configure headers and column widths
    hdr_cells = table.rows[0].cells
    for i, col_def in enumerate(config_cols):
        table.columns[i].width = Inches(col_def['width'])
        hdr_cells[i].text = col_def['header']
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Helper for safe formatting
    def safe_format(value, fmt):
        if isinstance(value, (int, float)):
            return fmt.format(value)
        return str(value or 'N/A')

    # Add financial data rows
    for item in financial_data:
        row_cells = table.add_row().cells
        
        # Prepare formatted data strings
        amount_str = safe_format(item.get('amount'), "{:,.2f}")
        score_str = safe_format(item.get('anomaly_score'), "{:.4f}")
        
        # Map data to cells based on defined column order
        data_map = [
            item.get('vendor', 'N/A'),
            amount_str,
            item.get('currency', 'N/A'),
            score_str
        ]
        
        for i, (text, col_def) in enumerate(zip(data_map, config_cols)):
            row_cells[i].text = text
            # Apply alignment based on configuration
            row_cells[i].paragraphs[0].alignment = col_def['alignment']


def _add_anomaly_summary(document, anomalies):
    """Adds a detailed section for detected anomalies."""
    document.add_heading('2. Detected Anomalies', level=2)

    if not anomalies:
        document.add_paragraph("No specific high-priority anomalies were flagged in the detailed analysis.")
        return

    document.add_paragraph("The following transactions exceeded predefined anomaly thresholds:")
    
    # Use an unnumbered list
    for i, anomaly in enumerate(anomalies):
        amount_raw = anomaly.get('amount', 'N/A')
        amount_str = f"{amount_raw:,.2f}" if isinstance(amount_raw, (int, float)) else str(amount_raw)
        details = anomaly.get('details', 'General anomaly.') 
        
        p = document.add_paragraph(style='List Bullet')
        p.add_run(f"Anomaly {i+1} (").bold = True
        p.add_run(f"{anomaly.get('id', 'No ID')}").font.size = Pt(9) # Hallucinated transaction ID for reference
        p.add_run("): ").bold = True
        p.add_run(f"Amount: {amount_str} / Details: {details}")


def _add_overall_assessment(document, is_suspicious):
    """Adds the final summary and assessment section."""
    document.add_heading('3. Overall Assessment', level=2)

    if is_suspicious:
        p = document.add_paragraph()
        p.add_run("STATUS: HIGH RISK / SUSPICIOUS ACTIVITY DETECTED").bold = True
        p.add_run("\nThis document triggered critical alerts based on applied heuristics. Further manual review is highly recommended.")
        document.add_page_break()
    else:
        document.add_paragraph("STATUS: LOW RISK / CLEAN - All monitored metrics fall within acceptable parameters.")


def generate_report(pdf_info, output_file):
    """
    Generate a comprehensive financial review report using python-docx.

    Args:
        pdf_info (dict): Dictionary containing financial data, anomalies, and status.
        output_file (str): Path to the output DOCX file.
    """
    document = None
    try:
        document = Document()
        
        _add_report_title(document, "Financial Document Review Report")
        
        _add_financial_data_table(document, pdf_info.get('financial_data', []))

        _add_anomaly_summary(document, pdf_info.get('anomalies', []))

        _add_overall_assessment(document, pdf_info.get('suspicious', False))
        
        # Save report
        document.save(output_file)
        logging.info(f"[REPORT] Report successfully generated: {output_file}")
        
    except Exception as e:
        logging.error(f"[REPORT_SAVE_ERROR] Failed to generate or save DOCX file '{output_file}': {e}")
        # Re-raise the exception to inform upstream processes
        raise


if __name__ == '__main__':
    # Example usage for testing and demonstration
    try:
        pdf_info = {
            'financial_data': [
                {'vendor': 'Vendor Corp 1', 'amount': 100.55, 'currency': 'USD', 'anomaly_score': 0.005},
                {'vendor': 'Vendor Inc 2 (High Score)', 'amount': 200000.75, 'currency': 'EUR', 'anomaly_score': 0.8256}
            ],
            'anomalies': [
                {'id': 'TXN1023', 'amount': 500.0, 'details': 'Transaction deviates 4 sigma from vendor baseline.'},
                {'id': 'TXN1024', 'amount': 150000.0, 'details': 'Large single wire transfer detected.'}
            ],
            'suspicious': True
        }
        output_file = 'report_enhanced.docx'
        generate_report(pdf_info, output_file)

        # Test case: No data
        generate_report({'financial_data': [], 'anomalies': [], 'suspicious': False}, 'report_empty.docx')

    except Exception as e:
        logging.error(f"[MAIN_EXECUTION] Critical error during report process: {e}")
        traceback.print_exc()