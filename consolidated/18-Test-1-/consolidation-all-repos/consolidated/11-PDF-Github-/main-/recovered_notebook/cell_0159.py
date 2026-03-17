import logging
from typing import Dict, Any, List
import traceback

from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# --- Type Definitions ---
ReportData = Dict[str, Any]

class WebScrapingUtilities:
    """Collection of static utilities related to content access policies.
    Note: This specifically checks HTML meta tags, not robots.txt files.
    """

    @staticmethod
    def check_html_robots_meta(html_soup: BeautifulSoup) -> bool:
        """
        Checks if HTML content contains a 'noindex' or 'nofollow' robots meta tag.
        
        Args:
            html_soup (BeautifulSoup): The parsed HTML content.

        Returns:
            bool: True if allowed (no restrictive meta tag found), False otherwise.
        """
        if not html_soup:
            logging.warning(f"[META_ROBOTS] Empty BeautifulSoup object provided, assuming allowed.")
            return True
            
        try:
            allowed_tag = html_soup.find('meta', {'name': 'robots'})
            if allowed_tag:
                content = allowed_tag.get('content', '').lower()
                if 'noindex' in content or 'nofollow' in content:
                    logging.info(f"[META_ROBOTS] Disallowed by meta tag content: {content}.")
                    return False
            return True
        except Exception:
            logging.warning(f"[META_ROBOTS] Could not parse robots meta tag due to error, assuming allowed.")
            return True


def simulate_content(is_corrupted: bool = False) -> str:
    """Simulates the content extracted from a document (e.g., PDF).

    Args:
        is_corrupted (bool): Whether to simulate a data corruption scenario.
    """
    if not is_corrupted:
        return "Clean PDF Content\nPayment to Vendor A: $100.00\nPayment to Vendor B: $200.00\n"
    else:
        # Corrupted example: missing decimal place
        return "Corrupted PDF Content\nPayment to Vendor A: $100.00\nPayment to Vendor B: $20000"


class DocumentGenerator:
    """Handles generating output reports (e.g., DOCX) from structured analysis data."""

    def __init__(self, logger: logging.Logger):
        self.logger = logger

    def generate_docx_report(self, pdf_info: ReportData, output_file: str) -> None:
        """
        Generates a DOCX report based on the provided analysis results.
        
        Args:
        pdf_info (ReportData): The structured analysis information.
        output_file (str): The output file path.
        """
        document = Document()
        document.add_heading("Document Processing Analysis Report", level=1)
        
        financial_data: List[Dict[str, str]] = pdf_info.get('financial_data', [])
        anomalies: List[Dict[str, str]] = pdf_info.get('anomalies', [])
        is_suspicious: bool = pdf_info.get('suspicious', False)

        # 1. Core Financial Data
        document.add_heading(f"1. Processed Transactions ({len(financial_data)} Items)", level=2)
        if financial_data:
            for item in financial_data:
                document.add_paragraph(f"Amount: {item.get('amount', 'N/A')}")
        else:
            document.add_paragraph("No core financial data extracted.")

        # 2. Anomaly Detection
        document.add_heading(f"2. Anomalies Detected ({len(anomalies)})", level=2)
        
        if anomalies:
            for anomaly in anomalies:
                document.add_paragraph(f"!!! SUSPICIOUS AMOUNT: {anomaly.get('amount', 'N/A')}", style='Strong')
        else:
            document.add_paragraph("No parsing or validation anomalies detected.")

        # 3. Overall Verdict
        document.add_heading("3. Final Verdict", level=2)
        if is_suspicious:
            document.add_paragraph("Document flagged as SUSPICIOUS due to data inconsistencies or large anomalies.", style='Intense Quote')
        else:
            document.add_paragraph("Document appears clean and validated.")

        try:
            document.save(output_file)
            self.logger.info(f"[REPORT] Report successfully generated: {output_file}")
        except Exception as e:
            self.logger.error(f"Failed to save document {output_file}: {e}")
            raise

def main() -> None:
    """
    The main function, demonstrating report generation.
    """
    logger = logging.getLogger(__name__)
    report_generator = DocumentGenerator(logger)

    try:
        # Simulate analysis results
        pdf_info_analysis = {
            'document_name': 'Vendor_Invoice_Q4',
            'financial_data': [
                {'amount': '100.00'},
                {'amount': '200.00'}
            ],
            'anomalies': [
                {'amount': '500.00', 'reason': '5x standard deviation'}
            ],
            'suspicious': True
        }
        
        output_file = 'analysis_report.docx'
        report_generator.generate_docx_report(pdf_info_analysis, output_file)
        
        # Example usage of utility
        # Note: Must supply valid BeautifulSoup content here, not just 'robots.txt' text
        dummy_html = "<html><head><meta name='robots' content='noindex'></head><body></body></html>"
        soup = BeautifulSoup(dummy_html, 'html.parser')
        is_allowed = WebScrapingUtilities.check_html_robots_meta(soup)
        logger.info(f"[DEMO] HTML meta check result (should be False): {is_allowed}")
        
        logger.info(f"[DEMO] Simulated clean content: {simulate_content(False)[:30]}...")

    except Exception as e:
        logger.error(f"[MAIN] Error in execution: {e}")
        traceback.print_exc()

if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        logging.error(f"A critical error occurred: {e}")
        traceback.print_exc()