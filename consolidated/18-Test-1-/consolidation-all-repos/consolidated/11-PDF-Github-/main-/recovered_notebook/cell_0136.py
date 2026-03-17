import requests
from bs4 import BeautifulSoup
import re
from docx import Document
from xml.etree import ElementTree as ET
import concurrent.futures

# Define the sitemap namespace
SITEMAP_NS = 'http://www.sitemaps.org/schemas/sitemap/0.9'

def fetch_sitemap(sitemap_url):
    """
    Fetch and parse the sitemap at the given URL.

    Args:
        sitemap_url (str): The URL of the sitemap to fetch.

    Returns:
        list: A list of URLs found in the sitemap.
    """
    try:
        # Increased timeout slightly for robustness
        response = requests.get(sitemap_url, timeout=20)
        response.raise_for_status()

        # Use ElementTree default parser
        root = ET.fromstring(response.content)
        
        # Register namespace for cleaner XPath searches, though findall works with prefix
        namespaces = {'s': SITEMAP_NS}

        # Search for 'loc' tags, handling both sitemap and sitemapindex structures
        urls = [loc.text for loc in root.findall('s:url/s:loc', namespaces) if loc.text] 
        
        # Handle Sitemaps Index (contains <sitemap> elements pointing to other sitemaps)
        if not urls:
             index_urls = [loc.text for loc in root.findall('s:sitemap/s:loc', namespaces) if loc.text] 
             # Note: Recursive fetching is omitted for simplicity, but the structure identifies index files.

        return urls

    except requests.exceptions.RequestException as e:
        print(f"[ERROR] Error fetching sitemap {sitemap_url}: {e}")
        return []

    except ET.ParseError as e:
        print(f"[ERROR] Error parsing sitemap {sitemap_url}: {e}")
        return []

    except Exception as e:
        print(f"[ERROR] An unexpected error occurred processing sitemap {sitemap_url}: {e}")
        return []

def scrape_page_for_info(url):
    """
    Scrape the given URL for email addresses, phone numbers, and forms.
    Returns results as a structured dictionary.
    
    Args:
        url (str): The URL to scrape.
        
    Returns:
        dict: Dictionary containing scrape results or error information.
    """
    result = {"url": url, "emails": [], "phones": [], "forms": [], "error": None}
    
    try:
        response = requests.get(url, timeout=15)
        response.raise_for_status()
        content_text = response.text
        soup = BeautifulSoup(response.content, 'html.parser')

        # 1. Emails
        result["emails"] = sorted(list(set(re.findall(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", content_text))))

        # 2. Phone Numbers (Improved regex for slight robustness)
        # Note: This regex is still highly localized, a full international standard requires major libraries.
        result["phones"] = sorted(list(set(re.findall(
            r"(?:\+?\d{1,3}[-.\s]?)?(?:\(?\d{2,4}\)?[-\.\s]?)?\d{3}[-\.\s]?\d{4,9}", 
            content_text
        ))))

        # 3. Forms
        form_actions = []
        for form in soup.find_all('form'):
            action = form.get('action', 'N/A')
            method = form.get('method', 'GET/POST')
            form_actions.append(f"Action: {action}, Method: {method}")
        result["forms"] = form_actions

    except requests.exceptions.RequestException as e:
        result["error"] = f"Request failed: {e}"

    except Exception as e:
        result["error"] = f"Unexpected scraping error: {e}"
        
    return result

def write_results_to_document(results, filename="scraping_results.docx"):
    """
    Sequentially writes the collected structured results into a Word document.
    """
    document = Document()
    document.add_heading("Web Scraping Report (Concurrent Fetch)", level=1)
    
    for result in results:
        url = result["url"]
        document.add_heading(f"Scraping Results for: {url}", level=2)

        if result["error"]:
             document.add_paragraph(f"Status: ERROR. {result['error']}").runs[0].font.color.rgb = requests.utils.docx_constants.RGBColor(0xFF, 0x00, 0x00)
             continue

        # Emails
        if result["emails"]:
            document.add_paragraph("Email Addresses:")
            for email in result["emails"]:
                document.add_paragraph(f"   - {email}", style='List Bullet')
        else:
            document.add_paragraph("Email Addresses: None found.")

        # Phones
        if result["phones"]:
            document.add_paragraph("Phone Numbers:")
            for phone in result["phones"]:
                document.add_paragraph(f"   - {phone}", style='List Bullet')
        else:
            document.add_paragraph("Phone Numbers: None found.")

        # Forms
        if result["forms"]:
            document.add_paragraph("Forms Found:")
            for form_info in result["forms"]:
                document.add_paragraph(f"   - {form_info}", style='List Bullet')
        else:
            document.add_paragraph("Forms Found: None found.")
        
        document.add_paragraph('---') # Separator

    document.save(filename)

def main():
    # Optimal workers usually depend on network bandwidth (IO-bound task)
    MAX_WORKERS = 16 
    sitemap_url = "https://example.com/sitemap.xml"
    print(f"Fetching sitemap from: {sitemap_url}")

    urls = fetch_sitemap(sitemap_url)
    
    if not urls:
        print("No URLs found or sitemap failed to parse. Exiting.")
        return
    
    print(f"Found {len(urls)} URLs. Starting concurrent scraping with {MAX_WORKERS} workers.")

    results = []
    
    # Use ThreadPoolExecutor for concurrent I/O operations
    with concurrent.futures.ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
        future_to_url = {executor.submit(scrape_page_for_info, url): url for url in urls}
        
        for i, future in enumerate(concurrent.futures.as_completed(future_to_url)):
            url = future_to_url[future]
            try:
                data = future.result()
                results.append(data)
                print(f"[{i+1}/{len(urls)}] Scraped {url}")
            except Exception as exc:
                results.append({"url": url, "error": f"Future exception: {exc}"})
                print(f"[{i+1}/{len(urls)}] Error processing {url}: {exc}")

    # Write results sequentially
    write_results_to_document(results)
    print("Scraping complete. Results saved to scraping_results.docx")

if __name__ == "__main__":
    # Patch for docx color constant availability if needed, though often unnecessary in recent versions
    try:
        from docx.shared import RGBColor
        requests.utils.docx_constants = __import__('docx.shared', fromlist=['RGBColor'])
    except:
        pass

    main()