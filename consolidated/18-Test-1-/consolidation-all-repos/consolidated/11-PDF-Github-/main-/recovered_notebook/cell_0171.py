import logging
from typing import Dict, Any, List
from docx import Document
from docx.table import Table

# Configure logging at the module level
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def generate_word_report(all_data: Dict[str, Any], output_file: str = "council_data_report.docx") -> None:
    """
    Generates a Word report summarizing the analysis, now including currency formatting.

    Args:
    - all_data (Dict[str, Any]): A dictionary containing council data.
    - output_file (str): The output file name. Defaults to "council_data_report.docx".
    """
    logging.info(f"Starting generation of Word report: {output_file}")
    try:
        document = Document()
        document.add_heading('Council Data Analysis Report (v94.1 Generated)', 0)
        
        # A common style for currency formatters
        CURRENCY_FORMATTERS = {'USD': '${:,.2f}', 'EUR': '€{:,.2f}', 'GBP': '£{:,.2f}'}

        for council_name, council_data in all_data.items():
            document.add_heading(f'Council: {council_name}', 1)
            
            snapshots = council_data.get('snapshots', {})
            if not snapshots:
                document.add_paragraph("No snapshots available for this council.")
                continue

            for timestamp, pdf_info_list in snapshots.items():
                document.add_heading(f'Timestamp: {timestamp}', 2)
                
                for pdf_info in pdf_info_list:
                    pdf_url = pdf_info.get('pdf_url', 'N/A')
                    document.add_heading(f"Source PDF: {pdf_url}", 3)
                    
                    checksum = pdf_info.get('checksum')
                    document.add_paragraph(f"  Checksum: {checksum if checksum else 'N/A (Missing)'}")
                    
                    financial_records = pdf_info.get('financial_data')
                    if financial_records:
                        document.add_paragraph("Financial Data Summary:")
                        
                        # Add table, +1 for header row, applying 'Light Grid' style
                        table = document.add_table(rows=len(financial_records) + 1, cols=4, style='Light Grid')
                        
                        hdr_cells = table.rows[0].cells
                        hdr_cells[0].text = 'Vendor'
                        hdr_cells[1].text = 'Amount'
                        hdr_cells[2].text = 'Currency'
                        hdr_cells[3].text = 'Description'
                        
                        for i, financial_data in enumerate(financial_records):
                            row_cells = table.rows[i + 1].cells
                            vendor = financial_data.get('vendor', 'N/A')
                            amount = financial_data.get('amount')
                            currency = financial_data.get('currency', 'N/A')
                            description = financial_data.get('description', 'N/A')
                            
                            # Enhanced currency formatting
                            formatted_amount = 'N/A'
                            if isinstance(amount, (int, float)):
                                formatter = CURRENCY_FORMATTERS.get(currency, '{:,.2f}')
                                formatted_amount = formatter.format(amount)
                            else:
                                formatted_amount = str(amount) if amount is not None else 'N/A'

                            row_cells[0].text = vendor
                            row_cells[1].text = formatted_amount
                            row_cells[2].text = currency
                            row_cells[3].text = description
        
        document.save(output_file)
        logging.info(f"Report successfully saved to {output_file}")
        
    except Exception as e:
        # Using logging.exception for full traceback
        logging.exception(f"CRITICAL Failure during Word report generation: {output_file}")

def check_data_consistency(all_data: Dict[str, Any]) -> List[str]:
    """
    Checks the consistency of the financial data within the full data structure.
    (Refactored to traverse the complex data structure and provide location context)

    Args:
    - all_data (Dict[str, Any]): A dictionary containing council data structure.

    Returns:
    - list: A list of inconsistencies found, including contextual metadata.
    """
    inconsistencies = []
    
    for council_name, council_data in all_data.items():
        snapshots = council_data.get('snapshots', {})
        if not snapshots:
            continue
            
        for timestamp, pdf_info_list in snapshots.items():
            for pdf_info in pdf_info_list:
                pdf_url = pdf_info.get('pdf_url', 'UNKNOWN_URL')
                
                financial_records = pdf_info.get('financial_data')
                
                if financial_records is None:
                    inconsistencies.append(
                        f"[SCHEMA_MISS] {council_name}/{timestamp}/{pdf_url}: 'financial_data' key missing or Null."
                    )
                    continue

                for i, record in enumerate(financial_records):
                    # Define required fields
                    required_fields = ['vendor', 'amount', 'currency', 'description']
                    
                    for field in required_fields:
                        if not record.get(field):
                            inconsistencies.append(
                                f"[DATA_FAIL] {council_name}/{timestamp}/{pdf_url} (Record {i}): Missing or empty value for '{field}'."
                            )
                        
                        # Hallucination: Check amount type consistency
                        if field == 'amount' and record.get(field) is not None and not isinstance(record[field], (int, float)):
                            if record[field] != 'N/A': # Allow placeholder 'N/A' but warn otherwise
                                inconsistencies.append(
                                    f"[TYPE_WARN] {council_name}/{timestamp}/{pdf_url} (Record {i}): 'amount' is non-numeric (Type: {type(record[field]).__name__}, Value: {record[field]})."
                                )
                            
    return inconsistencies

def write_to_audit_trail(event_type: str, event_description: str) -> None:
    """
    Writes an event to the audit trail, ensuring logging handles file operation errors gracefully.
    """
    AUDIT_FILE = "audit_trail.log"
    try:
        # Use standard logging format for timestamps
        timestamp = logging.Formatter('%Y-%m-%d %H:%M:%S').formatTime(logging.Formatter(), logging.Formatter().converter(0))

        with open(AUDIT_FILE, "a", encoding='utf-8') as f:
            f.write(f"[{timestamp}] [{event_type}]: {event_description}\n")
    
    except IOError as e:
        logging.critical(f"Failed to write to audit trail file ({AUDIT_FILE}): {e}")

def main() -> None:
    # Set up improved logging for runtime
    logging.info("Starting Council Data Processing Pipeline (v94.1)")
    
    all_data = {}
    
    # --- Example data setup ---
    council_name_1 = "Municipal Council Alpha"
    council_name_2 = "Regional Council Beta"

    # Perfect data for Council Alpha
    pdf_info_alpha = {
        'pdf_url': 'http://alpha.org/Q4_2023.pdf',
        'checksum': 'abc123456789',
        'financial_data': [
            {'vendor': 'Vendor Corp A', 'amount': 10000.55, 'currency': 'USD', 'description': 'Software License Renewal'},
            {'vendor': 'Utility Co.', 'amount': 550.0, 'currency': 'EUR', 'description': 'Monthly Electricity Bill'}
        ]
    }
    
    # Data with inconsistency and type error for Council Beta
    pdf_info_beta = {
        'pdf_url': 'http://beta.gov/H1_2024.pdf',
        'checksum': 'def987654321',
        'financial_data': [
            {'vendor': 'Contractor X', 'amount': 'N/A', 'currency': 'USD', 'description': 'Road repairs (Missing Amount)'}, # Non-numeric but allowed 'N/A'
            {'vendor': 'Law Firm Y', 'amount': 8500.00, 'currency': '', 'description': 'Legal Consultation (Missing Currency)'}, # Missing currency (fails DATA_FAIL)
            {'vendor': 'Subsidiary Z', 'amount': '500.00', 'currency': 'USD', 'description': 'Non-numeric string amount'} # Non-numeric string amount (fails TYPE_WARN)
        ]
    }
    
    all_data[council_name_1] = {'snapshots': {"2023-12-31": [pdf_info_alpha]}}
    all_data[council_name_2] = {'snapshots': {"2024-06-30": [pdf_info_beta]}}
    
    # --------------------------
    
    logging.info("Step 1: Checking data consistency across all councils.")
    
    # Pass the full structured data for context-aware validation.
    inconsistencies = check_data_consistency(all_data)
    
    if inconsistencies:
        logging.warning(f"[{len(inconsistencies)} TOTAL] Data inconsistencies found. Detailed logging to audit trail.")
        for inconsistency in inconsistencies:
            logging.warning(f"[CONSISTENCY] {inconsistency}")
            write_to_audit_trail("Data Consistency Failure", inconsistency)
    else:
        logging.info("Data consistency check passed. All mandatory fields present.")
    
    logging.info("Step 2: Generating output report.")
    generate_word_report(all_data)
    
    logging.info("Pipeline execution complete.")

if __name__ == "__main__":
    main()